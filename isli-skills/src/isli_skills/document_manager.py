import io
import os
import markdown
import httpx
import structlog
from typing import Any, Union
try:
    from weasyprint import HTML
    HAS_WEASYPRINT = True
except (ImportError, OSError):
    HAS_WEASYPRINT = False
    
from docx import Document
from openpyxl import Workbook

from .auth import create_internal_token

logger = structlog.get_logger()

# WORKSPACE_URL will be imported from config or env at runtime to avoid circular imports
WORKSPACE_URL = os.getenv("WORKSPACE_URL", "http://localhost:8300")

async def save_to_workspace(agent_id: str, path: str, content: bytes) -> dict:
    """Upload generated binary content to the agent's workspace."""
    token = create_internal_token("isli-skills", scopes=["workspace:write"], expires_minutes=5)
    headers = {"X-Internal-Auth": token}
    
    # workspace /upload expects multipart/form-data
    files = {"file": (os.path.basename(path), content)}
    data = {
        "agent_id": agent_id,
        "path": path,
        "scope": "agent"
    }
    
    async with httpx.AsyncClient() as client:
        try:
            resp = await client.post(f"{WORKSPACE_URL}/upload", data=data, files=files, headers=headers)
            resp.raise_for_status()
            return resp.json()
        except httpx.HTTPStatusError as e:
            logger.error("workspace.upload_failed", status=e.response.status_code, detail=e.response.text)
            raise
        except Exception as e:
            logger.error("workspace.connection_failed", error=str(e))
            raise

def generate_pdf(content: str) -> bytes:
    """Generate PDF from Markdown or HTML."""
    if not content.strip().startswith("<"):
        # Basic Markdown to HTML conversion
        html_body = markdown.markdown(content, extensions=['tables', 'fenced_code'])
        # Add basic styling for PDF
        html_content = f"""
        <html>
        <head>
            <style>
                body {{ font-family: sans-serif; line-height: 1.6; margin: 2cm; }}
                table {{ border-collapse: collapse; width: 100%; margin-bottom: 1em; }}
                th, td {{ border: 1px solid #ccc; padding: 8px; text-align: left; }}
                th {{ background-color: #f2f2f2; }}
                pre {{ background-color: #f4f4f4; padding: 10px; border-radius: 5px; }}
            </style>
        </head>
        <body>{html_body}</body>
        </html>
        """
    else:
        html_content = content
    
    out = io.BytesIO()
    HTML(string=html_content).write_pdf(out)
    return out.getvalue()

def generate_docx(content: str) -> bytes:
    """Generate DOCX from Markdown (simplified)."""
    doc = Document()
    
    # Very basic markdown parser for DOCX
    # Improvement: Use a library like htmldocx or pypandoc for better fidelity
    lines = content.split('\n')
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line[2:], 0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 2)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
            
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def generate_xlsx(data: Union[list[dict[str, Any]], list[list[Any]]]) -> bytes:
    """Generate XLSX from a list of dictionaries or list of lists."""
    wb = Workbook()
    ws = wb.active
    
    if not data:
        ws.append(["No data provided"])
    elif isinstance(data[0], dict):
        headers = list(data[0].keys())
        ws.append(headers)
        for item in data:
            ws.append([item.get(h) for h in headers])
    elif isinstance(data[0], list):
        for row in data:
            ws.append(row)
    else:
        ws.append([str(data)])
            
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()

async def handle_document_generation(format: str, content: Any, filename: str, agent_id: str) -> dict:
    """Main entry point for document generation skill."""
    logger.info("document.generate.start", format=format, filename=filename, agent_id=agent_id)
    
    try:
        if format.lower() == 'pdf':
            if not HAS_WEASYPRINT:
                raise RuntimeError("PDF generation failed: weasyprint dependencies (pango, cairo) are missing in this environment.")
            binary_content = generate_pdf(str(content))
        elif format.lower() in ['docx', 'doc']:
            binary_content = generate_docx(str(content))
        elif format.lower() in ['xlsx', 'xls']:
            binary_content = generate_xlsx(content)
        else:
            raise ValueError(f"Unsupported document format: {format}")
            
        result = await save_to_workspace(agent_id, filename, binary_content)
        logger.info("document.generate.success", path=result.get("path"))
        return result
    except Exception as e:
        logger.error("document.generate.failed", error=str(e))
        raise
