from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_border(cell, **kwargs):
    """Helper to set borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_xml = OxmlElement(f'w:{edge}')
        for key, value in kwargs.items():
            edge_xml.set(qn(f'w:{key}'), str(value))
        tcPr.append(edge_xml)


def add_heading(doc, text, level=1):
    """Add a styled heading."""
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = 'Calibri'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    return p


def add_paragraph(doc, text, bold=False, italic=False, font_size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(6)
    return p


def add_table(doc, headers, rows):
    """Add a styled table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Calibri'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1F4E78')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)

    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            for paragraph in row_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    return table


def main():
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Renter Marketing Execution Plan')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(31, 78, 120)
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Demo Document — June 2026')
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    doc.add_paragraph()
    add_paragraph(doc, 'Prepared for: Renter (Car Rental SaaS)', bold=True)
    add_paragraph(doc, 'Prepared by: Donna — Project Orchestrator', bold=True)
    doc.add_paragraph()

    # 1. Executive Summary
    add_heading(doc, '1. Executive Summary', level=1)
    add_paragraph(doc,
        'Renter is positioned to capture a meaningful share of the fast-growing car rental software market, projected to expand from $2.86B in 2025 at 10.69% CAGR. The global car rental market is valued at $207.10B in 2026 and is expected to reach $557.62B by 2033.')
    add_paragraph(doc,
        'Consumer behavior is shifting: 49% of travelers prefer flexible rental models over ownership, and 62% of market growth is driven by urban travel demand. This creates a strong tailwind for Renter’s vertically-focused, direct-booking SaaS platform.')
    add_paragraph(doc,
        'Key Objective: Launch and execute a 12-month marketing program that drives demo requests, trial sign-ups, and paid conversions while building Renter’s brand as the modern operating system for independent rental businesses.', bold=True)

    # 2. Product Positioning
    add_heading(doc, '2. Product Positioning', level=1)
    add_paragraph(doc,
        'For independent car rental operators and fleet owners who are losing margin to OTAs and struggling with outdated software, Renter is the vertically-focused car rental SaaS that combines direct booking ownership with modern fleet management.')
    add_paragraph(doc,
        'Unlike Turo or legacy systems, Renter gives operators full control of their brand, customer data, and revenue streams without enterprise complexity.')

    add_heading(doc, 'Value Pillars', level=2)
    add_paragraph(doc, 'Own Your Revenue — Direct booking engine with 0% commission on your own channel.')
    add_paragraph(doc, 'Fleet Intelligence — Real-time availability, dynamic pricing, and maintenance alerts.')
    add_paragraph(doc, 'Launch in Days, Not Months — Cloud-native setup with modern UX and mobile-ready operations.')

    # 3. Target Audience
    add_heading(doc, '3. Target Audience', level=1)
    headers = ['Segment', 'Profile', 'Primary Pain Point']
    rows = [
        ['Independent Fleet Owners', '5–50 vehicles', 'High OTA commissions, no direct booking channel'],
        ['Peer-to-Peer Hosts', 'Individual car owners on Turo/Getaround', 'Platform dependency, high fees, no customer ownership'],
        ['Traditional Rental Agencies', 'Brick-and-mortar operators', 'Legacy software, poor UX, no real-time analytics'],
        ['Corporate Mobility Managers', 'SMEs managing employee travel', 'No centralized billing or policy compliance']
    ]
    add_table(doc, headers, rows)

    # 4. Competitive Landscape
    add_heading(doc, '4. Competitive Landscape', level=1)
    headers = ['Competitor', 'Type', 'Pricing', 'Strength', 'Weakness']
    rows = [
        ['Turo', 'P2P Marketplace', '10–40% guest fee', 'Large user base', 'High take rates, no customer ownership'],
        ['Getaround', 'P2P / On-demand', '40% revenue share', 'Connected car tech', 'Very high fees, limited host control'],
        ['YouRenti', 'SaaS', '~€49/month', 'Transparent pricing', 'EU-centric, limited enterprise features'],
        ['APPROVEX', 'SaaS', 'Custom / tiered', 'Back-office strength', 'Higher cost, steeper learning curve'],
        ['Levy Fleets', 'Fleet SaaS', 'From $24/month', 'Affordable entry', 'Basic feature set at low tier'],
        ['Traditional TMCs', 'Legacy Software', 'High enterprise contracts', 'Deep features', 'Outdated UX, long implementation']
    ]
    add_table(doc, headers, rows)

    # 5. Pricing
    add_heading(doc, '5. Renter Pricing Snapshot (Live API Data)', level=1)
    headers = ['Plan', 'Price', 'Billing Cycle', 'Max Cars', 'Max Users', 'Max Customers', 'Key Features']
    rows = [
        ['Free Trial', '$0', 'Monthly', '3', '2', '20', 'Maintenance, Expenses, Reporting'],
        ['Basic', '$199', 'Monthly', '10', '2', '200', 'Maintenance, Expenses, Reporting'],
        ['Premium', '$299', 'Monthly', '25', '4', '500', '+ GPS Tracking, Parks Module'],
        ['Enterprise', '$499', 'Monthly', '100', '10', '5,000', '+ Full modules'],
        ['Basic Y', '$2,000', 'Yearly', '10', '2', '1,000', '+ API Access']
    ]
    add_table(doc, headers, rows)

    # 6. Marketing Channels
    add_heading(doc, '6. Marketing Channels & Tactics', level=1)
    add_heading(doc, 'Priority Channels', level=2)
    headers = ['Priority', 'Channel', 'Tactic', 'KPI']
    rows = [
        ['P1', 'Content / SEO', 'Target keywords: “car rental software,” “Turo alternative for hosts,” “fleet management SaaS,” “direct booking car rental”', 'Organic traffic, demo requests'],
        ['P1', 'Product-Led Growth', 'Self-serve free trial with in-app activation dashboards', 'Trial-to-paid conversion rate'],
        ['P1', 'Partnerships', 'Insurance, telematics, and automotive lending integrations', 'Partner-sourced ARR'],
        ['P2', 'Paid Social', 'LinkedIn ads targeting fleet managers and rental owners', 'CAC, ROAS'],
        ['P2', 'Community', 'Build “Renter Collective” for hosts and operators', 'NPS, referrals'],
        ['P3', 'Events', 'Local fleet trade shows and mobility tech conferences', 'Qualified leads per event']
    ]
    add_table(doc, headers, rows)

    add_heading(doc, 'Content Pillars', level=2)
    add_paragraph(doc, 'The Anti-Marketplace Playbook — Reduce dependence on Turo/OTAs.')
    add_paragraph(doc, 'Fleet Ops Deep Dives — Maintenance, pricing, seasonality.')
    add_paragraph(doc, 'Operator Spotlights — Case studies of successful rental businesses.')
    add_paragraph(doc, 'SaaS Comparison Hub — Honest comparisons (Renter vs. legacy vs. marketplaces).')

    # 7. Launch Timeline
    add_heading(doc, '7. Launch Timeline', level=1)
    headers = ['Phase', 'Duration', 'Focus']
    rows = [
        ['Phase 1: Foundation', 'Month 1–2', 'Finalize messaging, launch SEO content hub, activate PLG free trial flow'],
        ['Phase 2: Acquisition', 'Month 3–4', 'Paid social campaigns, partnership announcements, referral program launch'],
        ['Phase 3: Scale', 'Month 5–6', 'Community launch, case study series, explore reseller/affiliate channels'],
        ['Phase 4: Expand', 'Month 7–12', 'Enterprise tier pilot, API ecosystem, international localization']
    ]
    add_table(doc, headers, rows)

    # 8. KPIs
    add_heading(doc, '8. 6-Month Success Metrics', level=1)
    headers = ['Metric', 'Target']
    rows = [
        ['Monthly Demo Requests', '150+'],
        ['Trial Sign-ups', '500+'],
        ['Trial-to-Paid Conversion', '12%+'],
        ['Net New ARR', '$50K+'],
        ['Organic Traffic Growth', '200%'],
        ['Customer Acquisition Cost', '<$400'],
        ['Lifetime Value', '>$3,000'],
        ['LTV:CAC Ratio', '> 3:1']
    ]
    add_table(doc, headers, rows)

    # 9. Risks
    add_heading(doc, '9. Risk & Mitigation', level=1)
    headers = ['Risk', 'Mitigation']
    rows = [
        ['Turo/OTA policy changes', 'Emphasize data ownership and direct customer relationships from day one'],
        ['Legacy software inertia', 'Offer white-glove onboarding and data migration services'],
        ['Price sensitivity from small fleets', 'Transparent, tiered pricing with clear ROI calculator']
    ]
    add_table(doc, headers, rows)

    # 10. Execution Plan
    add_heading(doc, '10. Execution Plan Summary', level=1)
    steps = [
        'Refresh market & competitor/keyword research and save findings to workspace.',
        'Build SEO content hub / marketing site pages targeting core keywords and value pillars.',
        'Implement self-serve free trial & in-app onboarding flow (PLG motion).',
        'Create launch content pillars and operator spotlight/case study material.',
        'Launch partnership outreach to insurance, telematics, and automotive lending providers.',
        'Stand up the Renter Collective community channel and referral program mechanics.',
        'Configure KPI dashboard and automated reporting for demo requests, trials, conversions, CAC/LTV.',
        'Plan and set up paid social (LinkedIn) campaigns targeting fleet managers and rental owners.',
        'Synthesize all deliverables, verify baselines, and present final execution report.'
    ]
    for i, step in enumerate(steps, start=1):
        add_paragraph(doc, f'{i}. {step}')

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('End of document')
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    doc.save('Renter_Marketing_Demo.docx')


if __name__ == '__main__':
    main()
